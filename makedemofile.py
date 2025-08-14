import pandas as pd
from docx import Document

def make_demoexcel():
    # 构造示例数据
    data = [
        {"客户名称": "张三公司", "合同编号": "HT2025-01", "签署日期": "2025-08-14", "合作内容": "技术服务", "合作期限": "1年"},
        {"客户名称": "李四集团", "合同编号": "HT2025-02", "签署日期": "2025-08-15", "合作内容": "设备采购", "合作期限": "2年"},
        {"客户名称": "王五有限公司", "合同编号": "HT2025-03", "签署日期": "2025-08-16", "合作内容": "市场推广", "合作期限": "半年"},
    ]

    # 保存 Excel
    df = pd.DataFrame(data)
    df.to_excel("data.xlsx", index=False)

    print("✅ data.xlsx 已生成")


def make_demoword():
    doc = Document()

    doc.add_heading('合作协议', level=1)

    doc.add_paragraph("甲方（客户）：{{ 客户名称 }}")
    doc.add_paragraph("合同编号：{{ 合同编号 }}")
    doc.add_paragraph("签署日期：{{ 签署日期 }}")

    doc.add_paragraph("甲方与乙方本着平等互利的原则，达成如下合作条款：")
    doc.add_paragraph("1. 合作内容：{{ 合作内容 }}")
    doc.add_paragraph("2. 合作期限：{{ 合作期限 }}")

    doc.add_paragraph("（以下略）")

    doc.save("合同模板.docx")
    print("✅ 合同模板.docx 已生成")


if __name__ == '__main__':
    make_demoexcel()
    make_demoword()